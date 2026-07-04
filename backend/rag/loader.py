"""Загрузка корпуса: PDF (постранично) и DOCX (текст + таблицы).

Картинки (схемы флотации, регламенты в PNG) в MVP не индексируются — им нужен
отдельный OCR/vision-слой; это честно отмечено как будущее расширение.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".docx"}
_PLANT_HINTS = {
    "кгмк": "KGMK",
    "kgmk": "KGMK",
    "ноф": "NOF",
    "nof": "NOF",
    "тоф": "TOF",
    "tof": "TOF",
}


@dataclass(frozen=True)
class DocUnit:
    """Единица текста внутри файла: страница PDF или весь DOCX."""

    page: int | None
    text: str


def iter_files(root: str | Path) -> list[Path]:
    root_path = Path(root)
    if not root_path.exists():
        raise FileNotFoundError(f"Corpus path not found: {root_path}")
    return sorted(
        path
        for path in root_path.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
    )


def detect_plant_hint(path: Path) -> str | None:
    haystack = str(path).lower()
    for needle, code in _PLANT_HINTS.items():
        if needle in haystack:
            return code
    return None


def kind_of(path: Path) -> str:
    return path.suffix.lower().lstrip(".")


def load_units(path: Path) -> list[DocUnit]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    return []


def _load_pdf(path: Path) -> list[DocUnit]:
    from pypdf import PdfReader

    units: list[DocUnit] = []
    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - битая страница не должна валить весь файл
            text = ""
        if text.strip():
            units.append(DocUnit(page=index, text=_normalize(text)))
    return units


def _load_docx(path: Path) -> list[DocUnit]:
    from docx import Document

    document = Document(str(path))
    parts = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = _normalize("\n".join(parts))
    return [DocUnit(page=None, text=text)] if text.strip() else []


def _normalize(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line)
