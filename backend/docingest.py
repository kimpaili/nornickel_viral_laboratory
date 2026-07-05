"""Приём экспертных гипотез из DOCX и текстовых документов в корпус.

Жюри приносит гипотезы «мозгового штурма» в DOCX (нумерованный список, часто
внутри таблицы из одной колонки), а книги/регламенты — в PDF. Этот модуль:

* DOCX → список экспертных гипотез (`hypothesis(origin='expert')`);
* PDF/текст → файл в папку корпуса для последующей индексации RAG.

Ключевой принцип этапа: ничего не должно падать на их данных. Любой кривой вход
даёт понятное сообщение и пустой результат, а не 500-ю ошибку.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import BinaryIO


# Заголовки-«шапки», которые не являются гипотезами.
_HEADER_MARKERS = (
    "гипотез",
    "мозгов",
    "результат",
    "предложен",
    "список",
)

# Инференс модуля-рычага по ключевым словам формулировки. Модуль не обязателен —
# если не угадали, гипотеза остаётся без привязки и оценивается всеми модулями.
_MODULE_KEYWORDS: list[tuple[str, tuple[str, ...]]] = [
    ("regrind", ("доизмельч", "измельч", "футеровк", "мельниц", "раскрыт", "помол")),
    (
        "classification",
        (
            "классифик",
            "гидроцикл",
            "циклон",
            "грохот",
            "грохоч",
            "насадк",
            "спираль",
            "гранулометр",
            "классификатор",
            "зазор",
            "дроблен",
            "дробилк",
        ),
    ),
    ("fine_flotation", ("флотац", "реагент", "пенн", "фронт")),
]

_NUMBER_PREFIX = re.compile(r"^\s*(?:\d+[.)]\s*)+")


@dataclass(frozen=True)
class ParsedHypothesis:
    title: str
    module_code: str | None


def infer_module(text: str) -> str | None:
    lowered = text.lower()
    for module_code, keywords in _MODULE_KEYWORDS:
        if any(word in lowered for word in keywords):
            return module_code
    return None


def _clean_line(raw: str) -> str:
    text = raw.replace("\xa0", " ").strip()
    text = _NUMBER_PREFIX.sub("", text).strip()
    # Убираем маркеры списка.
    text = text.lstrip("•—-–·*").strip()
    return text


def _is_header(text: str) -> bool:
    if text.endswith(":") and len(text) < 80:
        lowered = text.lower()
        return any(marker in lowered for marker in _HEADER_MARKERS)
    return False


def parse_hypotheses_from_docx(
    file: BinaryIO,
    filename: str,
) -> tuple[list[ParsedHypothesis], list[str]]:
    """Достаёт гипотезы из абзацев и таблиц DOCX.

    Возвращает (список гипотез, список предупреждений). Дубли и слишком короткие
    строки отсекаются, повторный текст не создаёт вторую гипотезу.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - зависимость есть в requirements
        raise ValueError("python-docx не установлен") from exc

    warnings: list[str] = []
    try:
        document = Document(file)
    except Exception as exc:  # noqa: BLE001 - любой битый DOCX -> понятное сообщение
        raise ValueError(
            f"Не удалось прочитать DOCX «{filename}»: файл повреждён или это не .docx"
        ) from exc

    raw_lines: list[str] = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                raw_lines.extend(cell.text.split("\n"))

    parsed: list[ParsedHypothesis] = []
    seen: set[str] = set()
    for raw in raw_lines:
        text = _clean_line(raw)
        if len(text) < 6:
            continue
        if _is_header(text):
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        parsed.append(ParsedHypothesis(title=text, module_code=infer_module(text)))

    if not parsed:
        warnings.append(
            f"{filename}: не найдено ни одной гипотезы — проверьте, что это список формулировок"
        )
    return parsed, warnings
